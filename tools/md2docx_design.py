# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\李一昂\Documents\Copywriting creat\docs\superpowers\specs\2026-08-04-badminton-copywriting-generator-design.md"
OUT = r"C:\Users\李一昂\Desktop\羽毛球甜妹文案生成器-设计文档.docx"

CN_FONT = "微软雅黑"
MONO = "Consolas"

def set_cn_font(run, name=CN_FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)

def shade_paragraph(p, fill):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)

def add_runs_with_inline(p, text, base_size=11, mono=False):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            set_cn_font(r, MONO if mono else CN_FONT)
            r.font.size = Pt(base_size)
        tok = m.group(0)
        if tok.startswith("**"):
            r = p.add_run(tok[2:-2])
            r.bold = True
            set_cn_font(r, MONO if mono else CN_FONT)
            r.font.size = Pt(base_size)
        else:
            r = p.add_run(tok[1:-1])
            set_cn_font(r, MONO)
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_cn_font(r, MONO if mono else CN_FONT)
        r.font.size = Pt(base_size)

with open(SRC, encoding="utf-8-sig") as f:
    lines = f.read().splitlines()

doc = Document()
style = doc.styles["Normal"]
style.font.size = Pt(11)
style.font.name = CN_FONT
style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

i = 0
in_code = False
code_buf = []

def flush_code():
    global code_buf
    for cl in code_buf:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run(cl if cl else " ")
        set_cn_font(r, MONO)
        r.font.size = Pt(9.5)
        shade_paragraph(p, "F2F4F7")
    code_buf = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    if stripped.startswith("```"):
        if in_code:
            flush_code()
            in_code = False
        else:
            in_code = True
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue
    if not stripped:
        i += 1
        continue
    if stripped.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(stripped[2:])
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = RGBColor(0x1A, 0x7F, 0x64)
        set_cn_font(r)
        i += 1
        continue
    if stripped.startswith("### "):
        p = doc.add_heading(level=2)
        r = p.add_run(stripped[4:])
        set_cn_font(r)
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x1A, 0x7F, 0x64)
        i += 1
        continue
    if stripped.startswith("## "):
        p = doc.add_heading(level=1)
        r = p.add_run(stripped[3:])
        set_cn_font(r)
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(0x1A, 0x7F, 0x64)
        i += 1
        continue
    if stripped.startswith("|"):
        tbl_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            tbl_lines.append(lines[i].strip())
            i += 1
        rows = []
        for tl in tbl_lines:
            if re.match(r"^\|[\s:\-|]+\|$", tl):
                continue
            cells = [c.strip() for c in tl.strip("|").split("|")]
            rows.append(cells)
        if rows:
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = table.rows[ri].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    txt = row[ci] if ci < len(row) else ""
                    add_runs_with_inline(p, txt, base_size=10.5)
                    if ri == 0:
                        for r in p.runs:
                            r.bold = True
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        continue
    m = re.match(r"^[-*]\s+(.*)$", stripped)
    if m:
        p = doc.add_paragraph(style="List Bullet")
        add_runs_with_inline(p, m.group(1))
        i += 1
        continue
    m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_runs_with_inline(p, m.group(2))
        i += 1
        continue
    p = doc.add_paragraph()
    add_runs_with_inline(p, stripped)
    i += 1

doc.save(OUT)
print("saved:", OUT)
